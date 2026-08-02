from __future__ import annotations

from io import BytesIO
from datetime import date, datetime

import pandas as pd
from docx import Document


def calculate_age(value: str) -> str:
    if not value:
        return ""
    try:
        born = pd.to_datetime(value).date()
        today = date.today()
        return str(today.year - born.year - ((today.month, today.day) < (born.month, born.day)))
    except Exception:
        return ""


def safe_date(value: str):
    try:
        return pd.to_datetime(value).date()
    except Exception:
        return date.today()


def player_label(row: pd.Series) -> str:
    club = row.get("club", "") or "Sin club"
    return f"{row.get('nombre', '')} — {club}"


def build_final_report_docx(player: dict, report: dict) -> bytes:
    document = Document()
    document.add_heading(player.get("nombre", "Informe de jugador"), 0)

    details = [
        ("Club", player.get("club", "")),
        ("Competición", player.get("competicion", "")),
        ("Posición", player.get("posicion_principal", "")),
        ("Pie", player.get("pie", "")),
        ("Altura", player.get("altura", "")),
        ("Fin de contrato", player.get("fin_contrato", "")),
    ]

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(details):
        cell = table.rows[index // 3].cells[index % 3]
        cell.text = f"{label}\n{value}"

    sections = [
        ("Observaciones", "observaciones"),
        ("Perfil físico", "perfil_fisico"),
        ("Perfil táctico", "perfil_tactico"),
        ("Perfil técnico", "perfil_tecnico"),
        ("Perfil mental", "perfil_mental"),
        ("Fases del juego", "fases_juego"),
        ("Áreas consolidadas", "areas_consolidadas"),
        ("Áreas de mejora", "areas_mejora"),
        ("Conclusión", "conclusion"),
        ("Características determinantes", "caracteristicas_determinantes"),
        ("Enlaces de video", "enlaces_video"),
    ]

    for title, field in sections:
        document.add_heading(title, level=1)
        document.add_paragraph(report.get(field, "") or "—")

    document.add_paragraph(
        f"Autor: {report.get('autor', '')} | Estado: {report.get('estado', '')}"
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
